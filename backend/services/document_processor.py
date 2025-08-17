"""
Document processing service for extracting text and metadata from various file types
"""

import os
import logging
import mimetypes
from typing import Dict, Any, Optional, Union
import tempfile
import aiofiles
from pathlib import Path

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import pandas as pd
from PIL import Image
import magic

from .advanced_document_processor import AdvancedDocumentProcessor

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing various document types"""
    
    def __init__(self):
        self.advanced_processor = AdvancedDocumentProcessor()
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'application/msword': self._process_doc,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'application/vnd.ms-excel': self._process_excel,
            'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': self._process_excel,
            'text/plain': self._process_text,
            'text/csv': self._process_csv,
            'image/jpeg': self._process_image,
            'image/png': self._process_image,
        }
    
    async def process_document(self, file_content: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """
        Process document and extract content and metadata
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            mime_type: MIME type of the file
        
        Returns:
            Dictionary containing extracted text, metadata, and processing info
        """
        try:
            # Validate file type
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {mime_type}")
            
            # Use advanced processing for PDFs
            if self.advanced_processor.is_pdf_file(mime_type):
                logger.info(f"Using advanced AI pipeline for PDF: {filename}")
                result = await self.advanced_processor.process_pdf_document(file_content, filename)
                
                # Add common metadata
                result.update({
                    'filename': filename,
                    'mime_type': mime_type,
                    'file_size': len(file_content)
                })
                
                return result
            
            # Use basic processing for other file types
            logger.info(f"Using basic processing for: {filename}")
            
            # Create temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix=Path(filename).suffix) as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                # Process based on file type
                processor = self.supported_types[mime_type]
                result = await processor(temp_file_path, filename)
                
                # Add common metadata
                result.update({
                    'filename': filename,
                    'mime_type': mime_type,
                    'file_size': len(file_content),
                    'processing_status': 'success'
                })
                
                logger.info(f"Successfully processed document: {filename}")
                return result
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception as e:
                    logger.warning(f"Failed to delete temp file {temp_file_path}: {e}")
                    
        except Exception as e:
            logger.error(f"Document processing failed for {filename}: {e}")
            return {
                'filename': filename,
                'mime_type': mime_type,
                'file_size': len(file_content),
                'processing_status': 'error',
                'error_message': str(e),
                'extracted_text': '',
                'metadata': {}
            }
    
    async def _process_pdf(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process PDF document"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text_content.append(page_text)
                
                extracted_text = '\n'.join(text_content)
                
                # Extract metadata
                metadata = {
                    'num_pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else '',
                    'creator': pdf_reader.metadata.get('/Creator', '') if pdf_reader.metadata else '',
                }
                
                return {
                    'extracted_text': extracted_text,
                    'metadata': metadata,
                    'word_count': len(extracted_text.split()),
                    'char_count': len(extracted_text)
                }
                
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise
    
    async def _process_docx(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process DOCX document"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            extracted_text = '\n'.join(text_content)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else '',
                'num_paragraphs': len(doc.paragraphs),
            }
            
            return {
                'extracted_text': extracted_text,
                'metadata': metadata,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise
    
    async def _process_doc(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process DOC document (legacy Word format)"""
        # For now, return basic info - would need python-docx2txt or similar for full processing
        return {
            'extracted_text': 'Legacy DOC format processing requires additional setup',
            'metadata': {'format': 'DOC (legacy)'},
            'word_count': 0,
            'char_count': 0
        }
    
    async def _process_excel(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process Excel document"""
        try:
            # Read Excel file
            df = pd.read_excel(file_path, sheet_name=None)  # Read all sheets
            
            # Extract text content from all sheets
            text_content = []
            total_rows = 0
            sheet_info = {}
            
            for sheet_name, sheet_df in df.items():
                sheet_text = sheet_df.to_string(index=False)
                text_content.append(f"Sheet: {sheet_name}\n{sheet_text}")
                
                sheet_info[sheet_name] = {
                    'rows': len(sheet_df),
                    'columns': len(sheet_df.columns),
                    'column_names': list(sheet_df.columns)
                }
                total_rows += len(sheet_df)
            
            extracted_text = '\n\n'.join(text_content)
            
            metadata = {
                'num_sheets': len(df),
                'total_rows': total_rows,
                'sheet_info': sheet_info
            }
            
            return {
                'extracted_text': extracted_text,
                'metadata': metadata,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"Excel processing failed: {e}")
            raise
    
    async def _process_text(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process plain text document"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                extracted_text = await file.read()
            
            lines = extracted_text.split('\n')
            
            metadata = {
                'num_lines': len(lines),
                'encoding': 'utf-8'
            }
            
            return {
                'extracted_text': extracted_text,
                'metadata': metadata,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            raise
    
    async def _process_csv(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process CSV document"""
        try:
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            extracted_text = df.to_string(index=False)
            
            metadata = {
                'num_rows': len(df),
                'num_columns': len(df.columns),
                'column_names': list(df.columns),
                'data_types': {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
            
            return {
                'extracted_text': extracted_text,
                'metadata': metadata,
                'word_count': len(extracted_text.split()),
                'char_count': len(extracted_text)
            }
            
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            raise
    
    async def _process_image(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Process image document"""
        try:
            with Image.open(file_path) as img:
                metadata = {
                    'format': img.format,
                    'mode': img.mode,
                    'size': img.size,
                    'width': img.width,
                    'height': img.height
                }
                
                # For now, images don't have extractable text (would need OCR)
                extracted_text = f"Image file: {filename} ({img.format}, {img.size[0]}x{img.size[1]})"
                
                return {
                    'extracted_text': extracted_text,
                    'metadata': metadata,
                    'word_count': len(extracted_text.split()),
                    'char_count': len(extracted_text)
                }
                
        except Exception as e:
            logger.error(f"Image processing failed: {e}")
            raise
    
    def get_file_type(self, file_content: bytes, filename: str) -> str:
        """Detect file MIME type from content and filename"""
        try:
            # Try to detect from file content
            mime_type = magic.from_buffer(file_content, mime=True)
            
            # Fallback to filename extension
            if mime_type == 'application/octet-stream':
                mime_type, _ = mimetypes.guess_type(filename)
            
            return mime_type or 'application/octet-stream'
            
        except Exception as e:
            logger.warning(f"MIME type detection failed: {e}")
            # Fallback to extension-based detection
            mime_type, _ = mimetypes.guess_type(filename)
            return mime_type or 'application/octet-stream'